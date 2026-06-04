"""Smoke test exportu do Wordu."""

from __future__ import annotations

from pathlib import Path

from app.core.models import (
    SECTION_KIND_BULLETS,
    SECTION_KIND_DEFINITIONS,
    SECTION_KIND_KEY_VALUE,
    SECTION_KIND_PARAGRAPH,
    SECTION_KIND_QA,
    SlideText,
    SourceFile,
    SourceKind,
    StudyMaterial,
    StudySection,
    Transcript,
    TranscriptSegment,
)
from app.core.word_export import (
    export_docx,
    suggested_output_filename,
    topic_folder_name,
)


def test_export_docx_creates_valid_file(tmp_path: Path) -> None:
    material = StudyMaterial(
        title="Testovací materiál",
        bullets=["První bod", "Druhý bod"],
        terms=[("pojem A", "definice A"), ("pojem B", "definice B")],
        examples=["Příklad 1"],
        further_study=["Dočti kapitolu 5"],
    )
    transcripts = [
        Transcript(
            source_label="Část 1",
            language="cs",
            duration_sec=120.0,
            text="Toto je testovací přepis přednášky.",
            segments=[
                TranscriptSegment(start=0.0, end=10.0, text="Toto je testovací"),
                TranscriptSegment(start=10.0, end=20.0, text="přepis přednášky."),
            ],
        )
    ]
    slides = [SlideText(source_label="Slidy 1", text="[Slide 1]\nNějaký text", slide_count=1)]
    sources = [
        SourceFile(path=tmp_path / "fake.mp3", kind=SourceKind.AUDIO_VIDEO, label="Část 1"),
    ]

    out = tmp_path / "test.docx"
    result = export_docx(
        output_path=out,
        material=material,
        transcripts=transcripts,
        slides=slides,
        sources=sources,
        user_prompt="Test prompt",
    )

    assert result == out
    assert out.exists()
    assert out.stat().st_size > 1000  # docx má nějaký rozumný objem


def test_suggested_filename_safe():
    material = StudyMaterial(title="Test / Materiál * 2026")
    name = suggested_output_filename(material)
    assert name.endswith(".docx")
    # žádné nelegální Windows znaky
    for forbidden in ("/", "\\", ":", "*", "?", '"', "<", ">", "|"):
        assert forbidden not in name


def test_suggested_filename_empty_title_fallback():
    material = StudyMaterial(title="")
    name = suggested_output_filename(material)
    assert name.endswith(".docx")
    assert len(name) > 5


def test_export_docx_includes_quiz_questions(tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    material = StudyMaterial(
        title="Hodina fyziky",
        bullets=["Newtonovy zákony"],
        quiz_questions=["Co říká první Newtonův zákon?", "Vysvětli setrvačnost."],
    )
    export_docx(output_path=out, material=material, transcripts=[], slides=[], sources=[], user_prompt=None)
    assert out.is_file()
    # Ověříme, že sekce s otázkami je v dokumentu
    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Otázky k procvičení a zkoušení" in full_text
    assert "první Newtonův zákon" in full_text


def test_topic_folder_name_basic():
    assert topic_folder_name(StudyMaterial(title="X", topic="Fyzika")) == "Fyzika"
    assert topic_folder_name(StudyMaterial(title="X", topic="Dějepis 20. století")) == "Dějepis 20 století"


def test_topic_folder_name_empty_returns_empty():
    assert topic_folder_name(StudyMaterial(title="X", topic="")) == ""
    assert topic_folder_name(StudyMaterial(title="X")) == ""


def test_topic_folder_name_strips_unsafe_chars():
    # Znaky nebezpečné pro Windows/macOS cesty se odstraní
    result = topic_folder_name(StudyMaterial(title="X", topic='Fyzika/Mechanika: úvod*?'))
    for forbidden in ("/", "\\", ":", "*", "?", '"', "<", ">", "|"):
        assert forbidden not in result
    assert "Fyzika" in result


def test_topic_folder_name_no_trailing_dot_or_space():
    # Windows nesnáší tečku/mezeru na konci názvu složky
    result = topic_folder_name(StudyMaterial(title="X", topic="Biologie. "))
    assert not result.endswith(".")
    assert not result.endswith(" ")


def test_topic_folder_name_length_capped():
    long_topic = "A" * 100
    assert len(topic_folder_name(StudyMaterial(title="X", topic=long_topic))) <= 40


def test_export_docx_renders_sections_format(tmp_path: Path) -> None:
    """Nový sekce-aware materiál: každý kind se musí v dokumentu projevit."""
    material = StudyMaterial(
        title="Sales schůzka — klient Novák",
        topic="Finance",
        sections=[
            StudySection(
                title="Úkoly pro mě",
                kind=SECTION_KIND_KEY_VALUE,
                items=[
                    ("Připravit srovnání tří fondů", "do pátku"),
                    ("Zavolat klientovi zpět", "neuvedeno"),
                ],
            ),
            StudySection(
                title="Profil klienta",
                kind=SECTION_KIND_KEY_VALUE,
                items=[("Věk", "42"), ("Děti", "2 (8, 11)")],
            ),
            StudySection(
                title="Termín další schůzky",
                kind=SECTION_KIND_PARAGRAPH,
                items=["Příští čtvrtek v 16:00, kancelář Praha 1."],
            ),
            StudySection(
                title="Další poznámky",
                kind=SECTION_KIND_BULLETS,
                items=["Klient se rozhoduje pomalu", "Manželka pozve příště"],
            ),
        ],
    )

    out = tmp_path / "sales.docx"
    export_docx(output_path=out, material=material, sources=[], user_prompt=None)

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Tituly všech sekcí
    assert "Úkoly pro mě" in full_text
    assert "Profil klienta" in full_text
    assert "Termín další schůzky" in full_text
    # Klíče i hodnoty key_value
    assert "Připravit srovnání tří fondů" in full_text
    assert "do pátku" in full_text
    # Paragraph
    assert "Příští čtvrtek" in full_text
    # Bullet
    assert "Klient se rozhoduje pomalu" in full_text


def test_export_docx_renders_qa_with_answers(tmp_path: Path) -> None:
    """QA sekce musí mít otázky tučně + vzorové odpovědi pod nimi."""
    material = StudyMaterial(
        title="Hodina dějepisu — opakování",
        sections=[
            StudySection(
                title="Otázky ke zkoušení",
                kind=SECTION_KIND_QA,
                items=[
                    ("Kdy začala 1. světová válka?", "V roce 1914."),
                    ("Vyjmenuj hlavní mocnosti Trojspolku.", "Německo, Rakousko-Uhersko, Itálie."),
                ],
            ),
        ],
    )
    out = tmp_path / "qa.docx"
    export_docx(output_path=out, material=material, sources=[], user_prompt=None)

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Kdy začala 1. světová válka?" in full_text
    assert "V roce 1914." in full_text


def test_export_docx_renders_definitions_with_bold_terms(tmp_path: Path) -> None:
    material = StudyMaterial(
        title="Fyzika — pojmy",
        sections=[
            StudySection(
                title="Klíčové pojmy",
                kind=SECTION_KIND_DEFINITIONS,
                items=[("síla", "vektorová veličina měřená v Newtonech")],
            )
        ],
    )
    out = tmp_path / "defs.docx"
    export_docx(output_path=out, material=material, sources=[], user_prompt=None)

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "síla" in full_text
    assert "vektorová veličina" in full_text


def test_export_docx_omits_transcript_and_slides(tmp_path: Path) -> None:
    """Plný přepis a obsah prezentací se do Wordu NEPÍŠOU — uživatel je má v .txt."""
    material = StudyMaterial(
        title="X",
        bullets=["Pouze tenhle bod"],
    )
    transcripts = [
        Transcript(
            source_label="Část 1",
            language="cs",
            duration_sec=120.0,
            text="UNIKATNI_PREPISOVY_TEXT_KTERY_BY_SE_NEMEL_OBJEVIT.",
            segments=[],
        )
    ]
    slides = [
        SlideText(
            source_label="Slidy",
            text="UNIKATNI_TEXT_ZE_SLIDU_KTERY_BY_SE_NEMEL_OBJEVIT.",
            slide_count=1,
        )
    ]
    out = tmp_path / "no_transcript.docx"
    export_docx(
        output_path=out,
        material=material,
        sources=[],
        user_prompt=None,
        transcripts=transcripts,
        slides=slides,
    )

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "UNIKATNI_PREPISOVY_TEXT" not in full_text
    assert "UNIKATNI_TEXT_ZE_SLIDU" not in full_text
    # Ale obsah materiálu tam být musí
    assert "Pouze tenhle bod" in full_text


def test_export_docx_renders_legacy_only_material(tmp_path: Path) -> None:
    """Materiál bez sections, jen s legacy bullets/terms — pořád vyrenderuje."""
    material = StudyMaterial(
        title="Legacy",
        bullets=["bod jedna", "bod dva"],
        terms=[("alfa", "definice alfy")],
        quiz_questions=["Co je alfa?"],
    )
    out = tmp_path / "legacy.docx"
    export_docx(output_path=out, material=material, sources=[], user_prompt=None)

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "bod jedna" in full_text
    assert "alfa" in full_text
    assert "Co je alfa?" in full_text


def test_export_docx_survives_control_characters(tmp_path: Path) -> None:
    """Control znaky (\\x00, \\x0b...) z Whisper/PDF nesmí shodit export —
    python-docx by jinak vyhodil ValueError a ztratila by se celá práce."""
    material = StudyMaterial(
        title="Mat\x00eriál\x0b s control znaky",
        bullets=["Bod s\x00 NULL", "Normální bod"],
        terms=[("po\x0cjem", "defi\x1fnice")],
        examples=["Příklad\x08"],
        further_study=["Zdroj\x0e"],
    )
    transcripts = [
        Transcript(
            source_label="Část\x00 1",
            language="cs",
            duration_sec=10.0,
            text="Přepis\x00 s control\x0b znakem.",
            segments=[TranscriptSegment(start=0.0, end=5.0, text="Přepis\x00 s control\x0b znakem.")],
        )
    ]
    out = tmp_path / "ctrl.docx"
    # Nesmí vyhodit ValueError
    result = export_docx(
        output_path=out,
        material=material,
        transcripts=transcripts,
        slides=[],
        sources=[],
        user_prompt="Prompt\x00 s NULL",
    )
    assert result.exists()
    assert out.stat().st_size > 1000
