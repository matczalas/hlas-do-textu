"""Generátor Wordového dokumentu (.docx) z StudyMaterial.

Struktura výstupu:
    Nadpis 0 — Titul materiálu
    Metadata (datum, zdroje, popis od uživatele)
    Nadpis 1 — sekce 1   (podle StudyMaterial.iter_sections())
        obsah podle kind (bullets / definitions / qa / key_value / paragraph)
    Nadpis 1 — sekce 2
        …

Plný přepis a obsah prezentací se sem už NEPÍŠOU — přepis je vedle Wordu
jako `prepis_*.txt` (auto-save v pipeline), slidy má uživatel ve zdrojovém
souboru. Word obsahuje výhradně AI výstup.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from loguru import logger

from app.core.models import (
    SECTION_KIND_BULLETS,
    SECTION_KIND_DEFINITIONS,
    SECTION_KIND_KEY_VALUE,
    SECTION_KIND_PARAGRAPH,
    SECTION_KIND_QA,
    SourceFile,
    StudyMaterial,
    StudySection,
)

# python-docx (resp. lxml) odmítne uložit XML s control znaky a vyhodí
# ValueError "All strings must be XML compatible". Whisper, Gemini i extrakce
# z PDF/PPTX občas takový znak protlačí (\x00, \x0b, \x0c). Bez sanitizace by
# spadl celý export NA KONCI pipeline — uživatel by ztratil hodiny přepisu.
# Povolujeme jen tab/newline/carriage-return.
_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean(text: str | None) -> str:
    """Odstraní XML-nekompatibilní control znaky. None → ''."""
    if not text:
        return ""
    return _CONTROL_CHARS_RE.sub("", text)


def export_docx(
    *,
    output_path: Path,
    material: StudyMaterial,
    sources: list[SourceFile] | None = None,
    user_prompt: str | None = None,
    # Zpětně kompatibilní parametry — pipeline je ještě posílá, ale Word
    # je úmyslně ignoruje (přepis je vedle v .txt, slidy v originálním souboru).
    transcripts: list | None = None,
    slides: list | None = None,
) -> Path:
    """Vytvoří .docx soubor na `output_path`. Vrací cestu k němu.

    Parametry `transcripts` a `slides` jsou zachované kvůli stávajícím volajícím
    (pipeline, regenerate, testy), ale do dokumentu se nepíšou.
    """
    del transcripts, slides  # ignorováno — viz docstring

    from docx import Document  # lazy

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_default_style(doc)

    # ----- Titul + metadata -----
    title = _clean(material.title).strip() or "Studijní materiál"
    doc.add_heading(title, level=0)
    _write_meta(doc, sources or [], user_prompt)

    # ----- Sekce -----
    any_rendered = False
    for section in material.iter_sections():
        if not section.items:
            continue
        _render_section(doc, section)
        any_rendered = True

    if not any_rendered:
        # AI nic nevyrobila a legacy pole jsou prázdná. Než aby uživatel dostal
        # prázdný dokument, vložíme informativní poznámku.
        doc.add_heading("Poznámka", level=1)
        doc.add_paragraph(
            "AI z přepisu nevytěžila žádný strukturovaný obsah. "
            "Otevři chat o dokumentu a zkus zadání upravit, nebo spusť "
            "regeneraci s jinou šablonou."
        )

    doc.save(str(output_path))
    logger.info("Uloženo: {} ({:.1f} KB)", output_path, output_path.stat().st_size / 1024)
    return output_path


def _setup_default_style(doc) -> None:
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _write_meta(doc, sources: list[SourceFile], user_prompt: str | None) -> None:
    """Krátká hlavička: datum, zdroje, zadání."""
    meta = doc.add_paragraph()
    meta.add_run(
        f"Vygenerováno: {datetime.now().strftime('%d. %m. %Y %H:%M')}\n"
    ).italic = True
    if sources:
        meta.add_run("Zdroje:\n").bold = True
        for src in sources:
            meta.add_run(f"  • {_clean(src.label)} ({_clean(src.path.name)})\n")
    if user_prompt:
        meta.add_run("\nZadání pro AI:\n").bold = True
        meta.add_run(_clean(user_prompt) + "\n")


def _render_section(doc, section: StudySection) -> None:
    """Vykreslí jednu sekci. Volí styl podle `kind`."""
    doc.add_heading(_clean(section.title), level=1)
    kind = section.kind
    if kind == SECTION_KIND_BULLETS:
        _render_bullets(doc, section.items)
    elif kind == SECTION_KIND_DEFINITIONS:
        _render_definitions(doc, section.items)
    elif kind == SECTION_KIND_QA:
        _render_qa(doc, section.items)
    elif kind == SECTION_KIND_KEY_VALUE:
        _render_key_value(doc, section.items)
    elif kind == SECTION_KIND_PARAGRAPH:
        _render_paragraphs(doc, section.items)
    else:
        # Defenzivně — pokud někdo přidal nový kind a zapomněl tady, vyrendrujeme bullets
        _render_bullets(doc, section.items)


def _render_bullets(doc, items: list) -> None:
    for item in items:
        text = _clean(str(item)).strip()
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _render_definitions(doc, items: list) -> None:
    for pair in items:
        term, definition = _split_pair(pair)
        if not term:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(_clean(term)).bold = True
        if definition:
            p.add_run(f" — {_clean(definition)}")


def _render_qa(doc, items: list) -> None:
    for pair in items:
        question, answer = _split_pair(pair)
        if not question:
            continue
        q_para = doc.add_paragraph(style="List Number")
        q_run = q_para.add_run(_clean(question))
        q_run.bold = True
        if answer:
            from docx.shared import Pt

            a_para = doc.add_paragraph(_clean(answer))
            a_para.paragraph_format.left_indent = Pt(20)
            for run in a_para.runs:
                run.italic = True


def _render_key_value(doc, items: list) -> None:
    for pair in items:
        key, value = _split_pair(pair)
        if not key and not value:
            continue
        p = doc.add_paragraph(style="List Bullet")
        if key:
            p.add_run(_clean(key)).bold = True
        if value:
            sep = ": " if key else ""
            p.add_run(f"{sep}{_clean(value)}")


def _render_paragraphs(doc, items: list) -> None:
    for item in items:
        text = _clean(str(item)).strip()
        if text:
            doc.add_paragraph(text)


def _split_pair(value) -> tuple[str, str]:
    """Vrátí (klíč, hodnota) z list/tuple/dict položky.

    Parser v routeru obvykle dodává list[(str, str)], ale chat a starý formát
    můžou sem propustit i dict nebo samotný string.
    """
    if isinstance(value, list | tuple):
        if len(value) == 0:
            return ("", "")
        if len(value) == 1:
            return (str(value[0]), "")
        return (str(value[0]), str(value[1]))
    if isinstance(value, dict):
        # Pořadí klíčů ne vždy spolehlivé — vezmeme to, co najdeme
        key = next(
            (
                str(value[k])
                for k in ("term", "pojem", "question", "otázka", "otazka", "key", "klíč", "klic", "name", "label")
                if value.get(k)
            ),
            "",
        )
        val = next(
            (
                str(value[k])
                for k in ("definition", "definice", "answer", "odpověď", "odpoved", "value", "hodnota", "text")
                if value.get(k)
            ),
            "",
        )
        return (key, val)
    return (str(value), "")


def safe_filename_part(text: str, *, fallback: str = "soubor", max_len: int = 60) -> str:
    """Sanitizuje text na bezpečnou část názvu souboru.

    Povolí písmena/čísla (i s diakritikou), mezery → pomlčky, ostatní znaky → _,
    ořízne na `max_len` a odstraní pomlčky/podtržítka na okrajích. Prázdný
    výsledek → `fallback`. Bez nelegálních znaků pro Windows/macOS.
    """
    safe = "".join(
        c if c.isalnum() or c in (" ", "-", "_") else "_" for c in text
    ).strip()
    safe = safe.replace(" ", "-")[:max_len].strip("-_")
    return safe or fallback


def suggested_output_filename(material: StudyMaterial) -> str:
    """Vrátí název souboru ve tvaru `Studijni-material_YYYY-MM-DD-HHMM.docx`."""
    safe_title = safe_filename_part(material.title, fallback="Studijni-material")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return f"{safe_title}_{timestamp}.docx"


def topic_folder_name(material: StudyMaterial) -> str:
    """Vrátí bezpečný název složky podle tématu, nebo "" když téma chybí.

    Téma navrhne AI (např. "Fyzika", "Dějepis"). Sanitizujeme na název složky
    bezpečný napříč Windows/macOS — bez `\\ / : * ? " < > |`, bez tečky na konci
    (Windows), max 40 znaků. Když je téma prázdné, vrátíme "" (export jde do
    kořenové výstupní složky jako dosud).
    """
    topic = (material.topic or "").strip()
    if not topic:
        return ""
    # Povolíme písmena, čísla, mezery, pomlčky, podtržítka; zbytek pryč
    safe = "".join(c if c.isalnum() or c in (" ", "-", "_") else " " for c in topic)
    safe = " ".join(safe.split())  # zkolabovat vícenásobné mezery
    safe = safe.strip(" .")[:40].strip()
    return safe
